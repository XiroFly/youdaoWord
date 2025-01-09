import json
import hashlib
from time import sleep
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests

# MD5哈希函数
def md5_hash(text):
    return hashlib.md5(text.encode('utf-8')).hexdigest()

# 生成sign参数
def generate_sign(text):
    v = "webdict"
    _ = "web"
    w = "Mk6hqtUp33DGGtoS63tTJbMUYjRrG1Lu"
    time = len(text + v) % 10
    r = text + v
    o = md5_hash(r)
    n = _ + text + str(time) + w + o
    sign = md5_hash(n)
    return sign, time

# Cookie转换函数
def parse_cookies(cookie_str):
    cookies = {}
    cookie_items = cookie_str.split('; ')
    for item in cookie_items:
        try:
            key, value = item.split('=', 1)
            cookies[key] = value
        except ValueError:
            continue  # 忽略格式不正确的cookie
    return cookies

# 请求数据体
def get_data(text):
    sign, time = generate_sign(text)
    data = {
        "q": text,
        "le": "en",
        "t": time,
        "client": "web",
        "sign": sign,
        "keyfrom": "webdict"
    }
    return data

# 获取单词详细内容
def get_special(text, url, headers, cookies,ifneedMore=False):
    try:
        response = requests.post(url, headers=headers, cookies=cookies, data=get_data(text))
        response.raise_for_status()
        if(ifneedMore):
            json_response = response.json().get("etym", {}).get("etyms", {}).get("zh", [])
            json_response_usphone = response.json().get("ec", {}).get("word", {}).get("usphone", "")
            json_response_trs = response.json().get("ec", {}).get("word", {}).get("trs", [])
            
            etymology = []
            trans=[]
            for tran in json_response_trs:
                trans.append({
                    tran.get("pos", "") : tran.get("tran", "")
                })
            for item in json_response:
                new_item = {
                 item.get("word", "").replace(":", ""): item.get("desc", "").replace(":", "") + " " + item.get("value", "").replace(":", "")
                }
                etymology.append(new_item)
            return {
                "word": text,
                "trans": trans,
                "usphone": json_response_usphone,
                "词源": etymology
            }
        else:
            json_response = response.json().get("etym", {}).get("etyms", {}).get("zh", [])
            filtered_items = []
            for item in json_response:
                new_item = {
                 item.get("word", "").replace(":", ""): item.get("desc", "").replace(":", "") + " " + item.get("value", "").replace(":", "")
                }
                filtered_items.append(new_item)
            return filtered_items
    except requests.RequestException as e:
        print(f"请求错误: {e}")
        return []

# 设置中文字体
def set_chinese_font(run, size):
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()  # 获取或添加 rPr 元素
    if rPr is not None:
        rFonts = rPr.get_or_add_rFonts()  # 获取或添加 rFonts 元素
        if rFonts is not None:
            rFonts.set(qn('w:eastAsia'), '宋体')

# 写入到DOCX
def add_word_to_docx(doc, word,  usphone, trans,etymology):
    word_paragraph = doc.add_paragraph()
    word_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_word = word_paragraph.add_run(word)
    run_word.bold = True
    set_chinese_font(run_word, 16)


    usphone_paragraph = doc.add_paragraph()
    usphone_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_usphone = usphone_paragraph.add_run("\t" + usphone)
    set_chinese_font(run_usphone, 12)
    
    for item in trans:
        for key, value in item.items():
            ety_paragraph = doc.add_paragraph()
            ety_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            ety_text = f"{key} {value}"
            run_ety = ety_paragraph.add_run("\t" + ety_text)
            set_chinese_font(run_ety, 12)

    for item in etymology:
        for key, value in item.items():
            ety_paragraph = doc.add_paragraph()
            ety_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            ety_text = f"{key} {value}"
            run_ety = ety_paragraph.add_run("\t" + ety_text)
            set_chinese_font(run_ety, 10)

# 从JSON中提取数据
def extract_data_from_json_by_id(json_data, target_book_id):
    items = json_data.get("data", {}).get("items", [])
    filtered_items = []
    for item in items:
        if item.get("b") == target_book_id:
            i=0
            if(i>500):
                i=i-500
                sleep(2)   
            word = item.get("c", "")
            usphone = json.loads(item.get("j", "{}")).get("usphone", "")
            trans = json.loads(item.get("j", "{}")).get("trs", [{}])[0].get("tran", "")
            etymology = get_special(word, special_url, headers, cookies)
            filtered_items.append({
                "word": word,
                "trans": trans,
                "usphone": usphone,
                "词源": etymology
            })
    return filtered_items

def extract_data_from_json(json_data):
    unique_words = set()
    un_words =set()
    items = json_data.get("data", {}).get("items", [])
    # for item in items:
    #     if item.get("d",{})=="en":
    #         word = item.get("c", "")
    #         unique_words.add((word))
    #     if item.get("b",{})=="fa524244d9404b0da7ae32854f1c8d56" or item.get("b",{})== "9dcc15a7c8d24d119ae2823bde02cab9":
    #         word = item.get("c", "")
    #         un_words.add((word))
    # unique_words=unique_words-un_words
    for item in items:
        if item.get("b",{})=="9dcc15a7c8d24d119ae2823bde02cab9":
            word = item.get("c", "")
            unique_words.add((word))
    
    doc = Document()
    for word in unique_words: 
        res = get_special(word, special_url, headers, cookies,ifneedMore=True)
        add_word_to_docx(doc, res["word"], res["usphone"], res["trans"], res["词源"])
        
    doc.save("youdao_陌生.docx")
    
    

# 配置HTTP请求参数
special_url = "https://dict.youdao.com/jsonapi_s?doctype=json&jsonversion=4"
headers = {
    "Accept": "application/json, text/plain, */*",
    "Accept-Encoding": "gzip, deflate, br, zstd",
    "Accept-Language": "zh-CN,zh;q=0.9",
    "Connection": "keep-alive",
    "Host": "dict.youdao.com",
    "Sec-Ch-Ua": '"Not/A)Brand";v="8", "Chromium";v="126", "Google Chrome";v="126"',
    "Sec-Ch-Ua-Mobile": "?0",
    "Sec-Ch-Ua-Platform": '"Windows"',
    "Sec-Fetch-Dest": "empty",
    "Sec-Fetch-Mode": "cors",
    "Sec-Fetch-Site": "same-site",
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36",
}
cookie = "OUTFOX_SEARCH_USER_ID_NCOO=1983874567.9361298; OUTFOX_SEARCH_USER_ID=1184831676@61.179.137.126; P_INFO=19853230309|1723648449|1|dict_logon|00&99|null&null&null#shd&370200#10#0|&0||19853230309; DICT_PERS=v2|urs-phone-web||DICT||web||-1||1723648449935||61.179.137.155||urs-phoneyd.31036000cd774217b@163.com||kGhH6yhHYl06ZOfU50LeS0O5hMlW0HPK0Tz6LTL0MeL0p4RHQuRfTZ0q4nMkG0LJu0T4hLqBO4lG0qyOfgzOfUG0; DICT_UT=urs-phoneyd.31036000cd774217b@163.com; _uetvid=4d4d4800540211efba4cc7557eb057fd; DICT_SESS=v2|W7t3Uaf_UmwukLeuk4UMRqSOLlWk4UEReLRHQBkLzf0Uf0f6yPM6Z0pykMQZhfpu0gLhLw4PMlMRJLnLJzhMUfRQBnfPZhfgz0; DICT_LOGIN=3||1728015247863"
cookies = parse_cookies(cookie)

# 处理并写入数据到DOCX
def write_to_docx(filtered_items):
    doc = Document()
    for item in filtered_items:
        add_word_to_docx(doc, item["word"], item["usphone"], item["trans"], item["词源"])
        sleep(10)
    doc.save("youdao_ALL.docx")
def get_special1(word):
    return get_special(word, special_url, headers, cookies,ifneedMore=True)
# 主函数
def main():
     # 从文件中读取JSON数据
    with open(r'C:\Users\fly\Desktop\新建文本文档.txt', 'r', encoding='utf-8') as file:
        json_data = json.load(file)
        #target_book_id = "9dcc15a7c8d24d119ae2823bde02cab9"
        #filtered_items = extract_data_from_json_by_id(json_data, target_book_id)
        #write_to_docx(filtered_items)
        extract_data_from_json(json_data)

if __name__ == "__main__":
    main()
