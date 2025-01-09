import requests
import json
import hashlib
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# MD5哈希函数
def md5_hash(text):
    return hashlib.md5(text.encode('utf-8')).hexdigest()

# 生成sign参数
def generate_sign(text):
    v = "webdict"
    _ = "web"
    w = "Mk6hqtUp33DGGtoS63tTJbMUYjRrG1Lu"
    time = len(text + v) % 10
    r = text + v
    o = md5_hash(r)
    n = _ + text + str(time) + w + o
    sign = md5_hash(n)
    return sign, time

# Cookie转换函数
def parse_cookies(cookie_str):
    cookies = {}
    cookie_items = cookie_str.split('; ')
    for item in cookie_items:
        key, value = item.split('=', 1)
        cookies[key] = value
    return cookies

# 请求数据体
def get_data(text):
    sign, time = generate_sign(text)
    data = {
        "q": text,
        "le": "en",
        "t": time,
        "client": "web",
        "sign": sign,
        "keyfrom": "webdict"
    }
    return data

# 获取单词详细内容
def get_Special(text):
    url = "https://dict.youdao.com/jsonapi_s?doctype=json&jsonversion=4"
    response1 = requests.post(url, headers=headers, cookies=cookies, data=get_data(text))
    if response1.status_code == 200:
        json_response = response1.json().get("etym", {}).get("etyms", {}).get("zh", [])
        filtered_items = []
        for item in json_response:
            new_item = {
                item.get("word").replace(":", ""): item.get("desc").replace(":", "") + " " + item.get("value").replace(":", "")
            }
            filtered_items.append(new_item)
        return filtered_items
    else:
        print("错误\n")
        return []

# 获取我的单词列表
def get_list():
    response = requests.get(url, headers=headers, params=params, cookies=cookies)
    if response.status_code == 200:
        data = response.json()
        items = data.get("data", {}).get("itemList", [])
        filtered_items = []
        for item in items:
            filtered_item = {
                "word": item.get("word"),
                "trans": item.get("trans"),
                "usphone": item.get("usphone"),
                "词源": get_Special(item.get("word"))
            }
            filtered_items.append(filtered_item)
        return filtered_items
    else:
        print(f"Failed to retrieve data. Status code: {response.status_code}")
        return []

# 设置中文字体
def set_chinese_font(run, size):
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()  # 获取或添加 rPr 元素
    if rPr is not None:
        rFonts = rPr.get_or_add_rFonts()  # 获取或添加 rFonts 元素
        if rFonts is not None:
            rFonts.set(qn('w:eastAsia'), '宋体')


# 写入到DOCX
def add_word_to_docx(doc, word, trans, usphone, etymology):
    word_paragraph = doc.add_paragraph()
    word_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_word = word_paragraph.add_run(word)
    run_word.bold = True
    set_chinese_font(run_word, 16)

    trans_paragraph = doc.add_paragraph()
    trans_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_trans = trans_paragraph.add_run("\t"+trans)
    set_chinese_font(run_trans, 12)

    usphone_paragraph = doc.add_paragraph()
    usphone_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_usphone = usphone_paragraph.add_run("\t"+usphone)
    set_chinese_font(run_usphone, 12)

    for item in etymology:
        for key, value in item.items():
            ety_paragraph = doc.add_paragraph()
            ety_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            ety_text = f"{key} {value}"
            run_ety = ety_paragraph.add_run("\t"+ety_text)
            set_chinese_font(run_ety, 10)

# 写入单词列表到DOCX
def write_to_docx():
    filtered_items = get_list()
    doc = Document()
    for item in filtered_items:
        add_word_to_docx(doc, item["word"], item["trans"], item["usphone"], item["词源"])
    doc.save("youdao_formatted.docx")

# 配置HTTP请求参数
url = "https://dict.youdao.com/wordbook/webapi/v2/word/list"
headers = {
    "Accept": "application/json, text/plain, */*",
    "Accept-Encoding": "gzip, deflate, br, zstd",
    "Accept-Language": "zh-CN,zh;q=0.9",
    "Connection": "keep-alive",
    "Host": "dict.youdao.com",
    "Sec-Ch-Ua": '"Not/A)Brand";v="8", "Chromium";v="126", "Google Chrome";v="126"',
    "Sec-Ch-Ua-Mobile": "?0",
    "Sec-Ch-Ua-Platform": '"Windows"',
    "Sec-Fetch-Dest": "empty",
    "Sec-Fetch-Mode": "cors",
    "Sec-Fetch-Site": "same-site",
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36",
}
params = {
    "limit": 48,
    "offset": 48,
    "sort": "time",
    "lanTo": "",
    "lanFrom": ""
}
cookie = "P_INFO=19853230309|1716722963|1|dict_logon|00&99|null&null&null#shd&370200#10#0|&0|null|19853230309; DICT_PERS=v2|urs-phone-web||DICT||web||-1||1716722964158||61.162.51.121||urs-phoneyd.31036000cd774217b@163.com||gyk4OAOfp40qS0HlGOfqy0JuhfY5RfwFRUW0MJBkfPK064PMw40LJ4RT4Of6ZO4qB0qyOMp4OLJZ0YWOM6FhMw40; DICT_UT=urs-phoneyd.31036000cd774217b@163.com; OUTFOX_SEARCH_USER_ID_NCOO=1983874567.9361298; OUTFOX_SEARCH_USER_ID=1184831676@61.179.137.126; _uetvid=4d4d4800540211efba4cc7557eb057fd; DICT_SESS=v2|R-tfUsHHaJBRMJBOLqy0Tyh4puhLgB0ez0HOMk46uRYmhHeKhMTZ0wFRLkW0Ley0QZ64YWP4pB0YEnHTZhMOM0Y5P4YEnMTF0; DICT_LOGIN=3||1723274421580"
cookies = parse_cookies(cookie)

# 执行写入操作
write_to_docx()
